import os
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import Depends, HTTPException, status
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from modules.connection_to_db.database import get_session
from modules.models.user import User
from modules.models.user_document import UserDocument
from modules.models.types import DocumentStatusEnum
from modules.schemas.document_schemas import (
    UserWithDocumentSummary,
    DocumentRejectRequest,
    DocumentStatus,
)
from modules.utils.admin_utils import get_current_admin

CONTRACT_CITY = "Великий Новгород"


def _resolve_contract_paths() -> tuple[Path, Path]:
    """Locate the DOCX template and ensure output directory exists.

    The service can be started from different working directories (e.g. docker
    image or local run). To avoid FileNotFoundError 500 errors, try several
    likely base paths until the template is found.
    """

    candidate_bases = [
        Path(__file__).resolve().parents[3],
        Path(__file__).resolve().parents[2],
        Path(__file__).resolve().parents[1],
    ]

    checked: list[str] = []

    for base in candidate_bases:
        template_path = base / "static" / "contract_template.docx"
        checked.append(str(template_path))

        if template_path.exists():
            generated_dir = base / "generated_contracts"
            os.makedirs(generated_dir, exist_ok=True)
            return template_path, generated_dir

    raise FileNotFoundError(
        "DOCX-шаблон не найден; проверены пути: " + ", ".join(checked)
    )


CONTRACT_DOCX_TEMPLATE_PATH, GENERATED_CONTRACTS_DIR = _resolve_contract_paths()


def _replace_in_paragraph(paragraph: Paragraph, values: dict[str, Any]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in values.items():
        placeholder = f"{{{key}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(val))

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_placeholders_in_docx(doc: DocxDocument, values: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            _replace_in_paragraph(paragraph, values)

        footer = section.footer
        for paragraph in footer.paragraphs:
            _replace_in_paragraph(paragraph, values)


def _week_word(n: int | None) -> str:
    if n is None:
        return ""
    n_abs = abs(n)
    last_two = n_abs % 100
    last = n_abs % 10
    if 11 <= last_two <= 14:
        return "недель"
    if last == 1:
        return "неделю"
    if 2 <= last <= 4:
        return "недели"
    return "недель"


def _render_contract_docx(user: User, doc: UserDocument) -> str:
    document = DocxDocument(CONTRACT_DOCX_TEMPLATE_PATH)

    document = DocxDocument(CONTRACT_DOCX_TEMPLATE_PATH)

    today_str = datetime.utcnow().strftime("%d.%m.%Y")
    week_word = _week_word(doc.weeks_count)

    values: dict[str, Any] = {
        "CITY": CONTRACT_CITY,
        "DATE": today_str,
        "FULL_NAME": doc.full_name or "",
        "ФИО": doc.full_name or "",
        "ADDRESS": doc.address or "",
        "PASSPORT": doc.passport or "",
        "PHONE": doc.phone or "",
        "EMAIL": user.email or "",
        "BANK_ACCOUNT": doc.bank_account or "-",
        "№_договора": doc.contract_number or "",
        "Номер_приложения": "1",
        "Серийный_номер_велик": doc.bike_serial or "",
        "Серийный_нормер_АКБ_1": doc.akb1_serial or "",
        "Серийный_нормер_АКБ_2": doc.akb2_serial or "",
        "Серийный_нормер_АКБ_3": doc.akb3_serial or "",
        "Сумма": doc.amount or "",
        "Сумма_пропись": doc.amount_text or "",
        "Кол_во_недель": str(doc.weeks_count) if doc.weeks_count is not None else "",
        "неделю": week_word,
        "Дата_заполнения": doc.filled_date or today_str,
        "Дат_конец_аренды": doc.end_date or "",
    }

    _replace_placeholders_in_docx(document, values)

    out_path = GENERATED_CONTRACTS_DIR / f"contract_user_{user.id}.docx"
    document.save(out_path)
    return str(out_path)


class AdminHandler:
    def __init__(
        self,
        db: Session = Depends(get_session),
        admin: User = Depends(get_current_admin),
    ):
        self.db = db
        self.admin = admin

    def list_users(self) -> list[UserWithDocumentSummary]:
        users = self.db.query(User).filter(User.role == "user").all()
        result: list[UserWithDocumentSummary] = []

        for u in users:
            doc = (
                self.db.query(UserDocument).filter(UserDocument.user_id == u.id).first()
            )
            result.append(
                UserWithDocumentSummary(
                    id=u.id,
                    email=u.email,
                    first_name=u.first_name,
                    last_name=u.last_name,
                    role=u.role,
                    document_status=DocumentStatus(doc.status) if doc else None,
                )
            )
        return result

    def get_user_document(self, user_id: int) -> UserDocument:
        doc = (
            self.db.query(UserDocument).filter(UserDocument.user_id == user_id).first()
        )
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )
        return doc

    def approve_document(self, user_id: int) -> UserDocument:
        doc = (
            self.db.query(UserDocument).filter(UserDocument.user_id == user_id).first()
        )
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )

        user = self.db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Пользователь не найден",
            )

        doc.status = DocumentStatusEnum.APPROVED
        doc.rejection_reason = None
        doc.contract_text = "Договор успешно сформирован"

        self.db.commit()
        self.db.refresh(doc)
        return doc

    def reject_document(
        self, user_id: int, body: DocumentRejectRequest
    ) -> UserDocument:
        doc = (
            self.db.query(UserDocument).filter(UserDocument.user_id == user_id).first()
        )
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )

        doc.status = DocumentStatusEnum.REJECTED
        doc.rejection_reason = body.reason
        doc.contract_text = None

        self.db.commit()
        self.db.refresh(doc)
        return doc

    def get_contract_docx_path(self, user_id: int) -> str:
        doc = (
            self.db.query(UserDocument).filter(UserDocument.user_id == user_id).first()
        )
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )

        if doc.status != DocumentStatusEnum.APPROVED:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Договор еще не одобрен",
            )

        user = self.db.query(User).filter(User.id == user_id).first()
        if not user:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Пользователь не найден",
            )

        try:
            return _render_contract_docx(user, doc)
        except FileNotFoundError as exc:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=str(exc),
            ) from exc
