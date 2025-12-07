import os
from datetime import datetime
from pathlib import Path
from typing import Any

from fastapi import Depends, HTTPException, status
from sqlalchemy.orm import Session
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from modules.connection_to_db.database import get_session
from modules.models.user import User
from modules.models.user_document import UserDocument
from modules.models.types import DocumentStatusEnum
from modules.schemas.document_schemas import UserDocumentUpdate
from modules.utils.jwt_utils import get_current_user


# ==== настройки DOCX (как в старом main, но локально в этом модуле) ====

BACKEND_DIR = Path(__file__).resolve().parents[3]

CONTRACT_DOCX_TEMPLATE_PATH = "static/contract_template.docx"
GENERATED_CONTRACTS_DIR = "generated_contracts"
CONTRACT_CITY = "Великий Новгород"

os.makedirs(GENERATED_CONTRACTS_DIR, exist_ok=True)


def _replace_in_paragraph(paragraph: Paragraph, values: dict[str, Any]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in values.items():
        placeholder = f"{{{key}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(val))

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_placeholders_in_docx(doc: DocxDocument, values: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            _replace_in_paragraph(paragraph, values)

        footer = section.footer
        for paragraph in footer.paragraphs:
            _replace_in_paragraph(paragraph, values)


def _week_word(n: int | None) -> str:
    if n is None:
        return ""
    n_abs = abs(n)
    last_two = n_abs % 100
    last = n_abs % 10
    if 11 <= last_two <= 14:
        return "недель"
    if last == 1:
        return "неделю"
    if 2 <= last <= 4:
        return "недели"
    return "недель"


def _render_contract_docx(user: User, doc: UserDocument) -> str:
    if not os.path.exists(CONTRACT_DOCX_TEMPLATE_PATH):
        raise FileNotFoundError("DOCX-шаблон не найден")

    document = DocxDocument(CONTRACT_DOCX_TEMPLATE_PATH)

    today_str = datetime.utcnow().strftime("%d.%m.%Y")
    week_word = _week_word(doc.weeks_count)

    values: dict[str, Any] = {
        "CITY": CONTRACT_CITY,
        "DATE": today_str,
        "FULL_NAME": doc.full_name or "",
        "ФИО": doc.full_name or "",
        "ADDRESS": doc.address or "",
        "PASSPORT": doc.passport or "",
        "PHONE": doc.phone or "",
        "EMAIL": user.email or "",
        "BANK_ACCOUNT": doc.bank_account or "-",
        "№_договора": doc.contract_number or "",
        "Номер_приложения": "1",
        "Серийный_номер_велик": doc.bike_serial or "",
        "Серийный_нормер_АКБ_1": doc.akb1_serial or "",
        "Серийный_нормер_АКБ_2": doc.akb2_serial or "",
        "Серийный_нормер_АКБ_3": doc.akb3_serial or "",
        "Сумма": doc.amount or "",
        "Сумма_пропись": doc.amount_text or "",
        "Кол_во_недель": str(doc.weeks_count) if doc.weeks_count is not None else "",
        "неделю": week_word,
        "Дата_заполнения": doc.filled_date or today_str,
        "Дат_конец_аренды": doc.end_date or "",
    }

    _replace_placeholders_in_docx(document, values)

    out_path = os.path.join(GENERATED_CONTRACTS_DIR, f"contract_user_{user.id}.docx")
    document.save(out_path)
    return out_path


class UserDocumentHandler:
    def __init__(
        self,
        db: Session = Depends(get_session),
        current_user: User = Depends(get_current_user),
    ):
        self.db = db
        self.user = current_user

    def _get_my_document(self) -> UserDocument | None:
        return (
            self.db.query(UserDocument)
            .filter(UserDocument.user_id == self.user.id)
            .first()
        )

    def get_my_document(self) -> UserDocument:
        doc = self._get_my_document()
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )
        return doc

    def upsert_my_document(self, data: UserDocumentUpdate) -> UserDocument:
        doc = self._get_my_document()

        if not doc:
            doc = UserDocument(
                user_id=self.user.id,
                full_name=data.full_name,
                address=data.address,
                passport=data.passport,
                phone=data.phone,
                bank_account=data.bank_account,
                contract_number=data.contract_number,
                bike_serial=data.bike_serial,
                akb1_serial=data.akb1_serial,
                akb2_serial=data.akb2_serial,
                akb3_serial=data.akb3_serial,
                amount=data.amount,
                amount_text=data.amount_text,
                weeks_count=data.weeks_count,
                filled_date=data.filled_date,
                end_date=data.end_date,
                status=DocumentStatusEnum.DRAFT,
            )
            self.db.add(doc)
        else:
            doc.full_name = data.full_name
            doc.address = data.address
            doc.passport = data.passport
            doc.phone = data.phone
            doc.bank_account = data.bank_account

            doc.contract_number = data.contract_number
            doc.bike_serial = data.bike_serial
            doc.akb1_serial = data.akb1_serial
            doc.akb2_serial = data.akb2_serial
            doc.akb3_serial = data.akb3_serial
            doc.amount = data.amount
            doc.amount_text = data.amount_text
            doc.weeks_count = data.weeks_count
            doc.filled_date = data.filled_date
            doc.end_date = data.end_date

            doc.status = DocumentStatusEnum.DRAFT
            doc.rejection_reason = None
            doc.contract_text = None

        self.db.commit()
        self.db.refresh(doc)
        return doc

    def submit_my_document(self) -> UserDocument:
        doc = self._get_my_document()
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Сначала заполните документ",
            )

        doc.status = DocumentStatusEnum.PENDING
        doc.rejection_reason = None

        self.db.commit()
        self.db.refresh(doc)
        return doc

    def get_my_contract_docx_path(self) -> str:
        doc = self._get_my_document()
        if not doc:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Документ не найден",
            )

        if doc.status != DocumentStatusEnum.APPROVED:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="Договор еще не одобрен",
            )

        return _render_contract_docx(self.user, doc)
