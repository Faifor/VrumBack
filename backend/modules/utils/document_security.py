from __future__ import annotations

from datetime import date, datetime
from functools import lru_cache
from pathlib import Path
from typing import Any, Mapping, TYPE_CHECKING

from cryptography.fernet import Fernet, InvalidToken
from docx import Document as DocxDocument
from docx.text.paragraph import Paragraph

from modules.utils.config import settings

if TYPE_CHECKING:
    from modules.models.user import User
    from modules.models.user_document import UserDocument


_PERSONAL_FIELDS = {
    "full_name",
    "inn",
    "registration_address",
    "residential_address",
    "passport",
    "phone",
    "bank_account",
}

# Numeric-like fields that should be converted back to integers after decryption.
_NUMERIC_FIELDS = {"inn", "passport", "bank_account", "amount"}

_DATE_FIELDS = {"filled_date", "end_date"}
_ENCRYPTED_DOCUMENT_FIELDS = {
    "contract_number",
    "bike_serial",
    "akb1_serial",
    "akb2_serial",
    "akb3_serial",
    "amount",
    "amount_text",
}
_DOCUMENT_FIELDS = _ENCRYPTED_DOCUMENT_FIELDS | _DATE_FIELDS

_SECURE_TEMPLATE_SUBDIR = "templates"
_SECURE_CONTRACTS_SUBDIR = "generated_contracts"
_ENCRYPTED_PREFIX = "enc:"
_CONTRACT_CITY = "Великий Новгород"


def _ensure_secure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    try:
        path.chmod(0o700)
    except PermissionError:
        # Best-effort hardening for environments where chmod is allowed.
        pass
    return path


def get_contract_template_path() -> Path:
    template_dir = _ensure_secure_dir(
        settings.SECURE_STORAGE_DIR / _SECURE_TEMPLATE_SUBDIR
    )
    return template_dir / settings.CONTRACT_TEMPLATE_FILENAME


def get_generated_contracts_dir() -> Path:
    return _ensure_secure_dir(settings.SECURE_STORAGE_DIR / _SECURE_CONTRACTS_SUBDIR)


def get_generated_contract_path(user_id: int) -> Path:
    return get_generated_contracts_dir() / f"contract_user_{user_id}.docx"


class SensitiveDataCipher:
    def __init__(self, key: str):
        try:
            self._fernet = Fernet(key.encode())
        except (TypeError, ValueError) as exc:
            raise ValueError("Invalid ENCRYPTION_KEY provided") from exc

    def encrypt(self, value: str | None) -> str | None:
        if value is None:
            return None
        if value.startswith(_ENCRYPTED_PREFIX):
            return value

        token = self._fernet.encrypt(value.encode())
        return f"{_ENCRYPTED_PREFIX}{token.decode()}"

    def decrypt(self, value: str | None) -> str | None:
        if value is None:
            return None
        if not value.startswith(_ENCRYPTED_PREFIX):
            return value

        token = value[len(_ENCRYPTED_PREFIX) :]
        try:
            return self._fernet.decrypt(token.encode()).decode()
        except InvalidToken:
            # If the token cannot be decrypted, return it unchanged to avoid data loss.
            return value


@lru_cache
def get_sensitive_data_cipher() -> SensitiveDataCipher:
    return SensitiveDataCipher(settings.ENCRYPTION_KEY)


def encrypt_document_fields(
    data: Mapping[str, Any],
    cipher: SensitiveDataCipher,
    allowed_fields: set[str] | None = None,
) -> dict[str, Any]:
    encrypted: dict[str, Any] = {}
    fields_to_encrypt = allowed_fields or (_PERSONAL_FIELDS | _DOCUMENT_FIELDS)
    for field in fields_to_encrypt:
        if field not in data:
            continue


        value = data.get(field)
        if field in _DATE_FIELDS:
            encrypted[field] = value
            continue

        encrypted[field] = cipher.encrypt(value if value is None else str(value))
    return encrypted

def _normalize_numeric(value: Any) -> Any:
    if value is None:
        return None
    try:
        return int(value)
    except (TypeError, ValueError):
        return value


def decrypt_user_fields(
    user: "User", cipher: SensitiveDataCipher
) -> dict[str, str | None]:
    decrypted: dict[str, str | None] = {}
    for field in _PERSONAL_FIELDS:
        raw_value = getattr(user, field)
        decrypted_value = cipher.decrypt(raw_value)

        if field in _NUMERIC_FIELDS:
            decrypted[field] = _normalize_numeric(decrypted_value)
        else:
            decrypted[field] = decrypted_value
    return decrypted


def decrypt_document_fields(
    doc: "UserDocument", cipher: SensitiveDataCipher
) -> dict[str, Any]:
    decrypted: dict[str, Any] = {}
    for field in _DOCUMENT_FIELDS:
        value = getattr(doc, field)
        if field in _DATE_FIELDS:
            decrypted[field] = value
        else:
            decrypted_value = cipher.decrypt(value)
            if field in _NUMERIC_FIELDS:
                decrypted[field] = _normalize_numeric(decrypted_value)
            else:
                decrypted[field] = decrypted_value
    return decrypted


def serialize_document_for_response(
    doc: "UserDocument | None", cipher: SensitiveDataCipher, user: "User | None" = None
) -> dict[str, Any]:
    if doc:
        doc.refresh_dates_and_status(update_active=False)
        user = user or doc.user
        doc_data = decrypt_document_fields(doc, cipher)
        filled_date = doc_data.get("filled_date")
        end_date = doc_data.get("end_date")
        doc_fields = {
            "id": doc.id,
            "contract_number": doc_data.get("contract_number"),
            "bike_serial": doc_data.get("bike_serial"),
            "akb1_serial": doc_data.get("akb1_serial"),
            "akb2_serial": doc_data.get("akb2_serial"),
            "akb3_serial": doc_data.get("akb3_serial"),
            "amount": doc_data.get("amount"),
            "amount_text": doc_data.get("amount_text"),
            "weeks_count": doc.weeks_count,
            "filled_date": _format_date_for_response(filled_date),
            "end_date": _format_date_for_response(end_date),
            "active": bool(doc.active),
            "contract_text": doc.contract_text,
        }
    else:
        doc_fields = {
            "id": None,
            "contract_number": None,
            "bike_serial": None,
            "akb1_serial": None,
            "akb2_serial": None,
            "akb3_serial": None,
            "amount": None,
            "amount_text": None,
            "weeks_count": None,
            "filled_date": None,
            "end_date": None,
            "active": False,
            "contract_text": None,
        }

    user = user or getattr(doc, "user", None)
    personal_data = decrypt_user_fields(user, cipher) if user else {}
    status = getattr(user, "status", None)
    rejection_reason = getattr(user, "rejection_reason", None)

    return {
        "full_name": personal_data.get("full_name"),
        "inn": personal_data.get("inn"),
        "registration_address": personal_data.get("registration_address"),
        "residential_address": personal_data.get("residential_address"),
        "passport": personal_data.get("passport"),
        "phone": personal_data.get("phone"),
        "bank_account": personal_data.get("bank_account"),
        "status": status.value if hasattr(status, "value") else str(status)
        if status
        else None,
        "rejection_reason": rejection_reason,
        **doc_fields,
    }

def _format_date_for_response(value: Any) -> str | None:
    if value is None:
        return None
    if isinstance(value, date):
        return value.isoformat()
    return str(value)


def _replace_in_paragraph(paragraph: Paragraph, values: dict[str, Any]) -> None:
    if not paragraph.runs:
        return

    full_text = "".join(run.text for run in paragraph.runs)
    new_text = full_text

    for key, val in values.items():
        placeholder = f"{{{key}}}"
        if placeholder in new_text:
            new_text = new_text.replace(placeholder, str(val))

    if new_text == full_text:
        return

    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""


def _replace_placeholders_in_docx(doc: DocxDocument, values: dict[str, Any]) -> None:
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, values)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_in_paragraph(paragraph, values)

    for section in doc.sections:
        header = section.header
        for paragraph in header.paragraphs:
            _replace_in_paragraph(paragraph, values)

        footer = section.footer
        for paragraph in footer.paragraphs:
            _replace_in_paragraph(paragraph, values)


def _week_word(n: int | None) -> str:
    if n is None:
        return ""
    n_abs = abs(n)
    last_two = n_abs % 100
    last = n_abs % 10
    if 11 <= last_two <= 14:
        return "недель"
    if last == 1:
        return "неделю"
    if 2 <= last <= 4:
        return "недели"
    return "недель"


def render_contract_docx(
    user: "User", doc: "UserDocument", decrypted_fields: Mapping[str, Any]
) -> str:
    template_path = get_contract_template_path()
    if not template_path.exists():
        raise FileNotFoundError(
            f"DOCX-шаблон не найден по пути: {template_path}. "
            "Поместите контрактный шаблон в SECURE_STORAGE_DIR/templates "
            "и обновите CONTRACT_TEMPLATE_FILENAME при необходимости."
        )

    document = DocxDocument(template_path)

    today_str = datetime.utcnow().strftime("%d.%m.%Y")
    week_word = _week_word(doc.weeks_count)

    filled_date_value = decrypted_fields.get("filled_date")
    end_date_value = decrypted_fields.get("end_date")
    filled_date_str = _format_date_human(filled_date_value) or today_str
    end_date_str = _format_date_human(end_date_value)

    values: dict[str, Any] = {
        "CITY": _CONTRACT_CITY,
        "DATE": today_str,
        "FULL_NAME": decrypted_fields.get("full_name") or "",
        "ФИО": decrypted_fields.get("full_name") or "",
        "ADDRESS": decrypted_fields.get("registration_address")
        or decrypted_fields.get("residential_address")
        or "",
        "REGISTRATION_ADDRESS": decrypted_fields.get("registration_address") or "",
        "RESIDENTIAL_ADDRESS": decrypted_fields.get("residential_address") or "",
        "PASSPORT": decrypted_fields.get("passport") or "",
        "PHONE": decrypted_fields.get("phone") or "",
        "INN": decrypted_fields.get("inn") or "",
        "EMAIL": user.email or "",
        "BANK_ACCOUNT": decrypted_fields.get("bank_account") or "-",
        "№_договора": decrypted_fields.get("contract_number") or "",
        "Номер_приложения": "1",
        "Серийный_номер_велик": decrypted_fields.get("bike_serial") or "",
        "Серийный_нормер_АКБ_1": decrypted_fields.get("akb1_serial") or "",
        "Серийный_нормер_АКБ_2": decrypted_fields.get("akb2_serial") or "",
        "Серийный_нормер_АКБ_3": decrypted_fields.get("akb3_serial") or "",
        "Сумма": decrypted_fields.get("amount") or "",
        "Сумма_пропись": decrypted_fields.get("amount_text") or "",
        "Кол_во_недель": str(doc.weeks_count) if doc.weeks_count is not None else "",
        "неделю": week_word,
        "Дата_аполнения": filled_date_str,
        "Дата_заполнения": filled_date_str,
        "Дат_конец_аренды": end_date_str,
    }

    _replace_placeholders_in_docx(document, values)

    out_path = get_generated_contract_path(user.id)
    document.save(out_path)
    return str(out_path)


def _format_date_human(value: Any) -> str:
    if not value:
        return ""
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    return str(value)